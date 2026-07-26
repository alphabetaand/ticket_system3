import os 
import logging
import psycopg2
from flask import Flask, request, jsonify, render_template_string, send_file
from flask_cors import CORS
from passlib.context import CryptContext
from docx import Document
from io import BytesIO
from datetime import datetime
from sqlalchemy import create_engine, Column, Integer, String, DateTime, func
from sqlalchemy.orm import declarative_base, sessionmaker
from sqlalchemy import and_, or_

# Configuration
QR_FOLDER = "qrcodes/"
DATABASE_URL = os.getenv("DATABASE_URL")
print("✅ Connexion à :", DATABASE_URL)

# Test connexion simple
try:
    conn = psycopg2.connect(DATABASE_URL)
    print("✅ Connexion PostgreSQL réussie !")
    conn.close()
except Exception as e:
    print("❌ Échec de la connexion :", e)

engine = create_engine(DATABASE_URL, echo=False)
Base = declarative_base()
SessionLocal = sessionmaker(bind=engine)
ADMIN_PASSWORD = "alphonse2000"
FLASK_PORT = 5000
MAX_HISTORY_ENTRIES = 50

# Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Sécurité
pwd_context = CryptContext(schemes=["pbkdf2_sha256"], default="pbkdf2_sha256", pbkdf2_sha256__default_rounds=30000)
ADMIN_PASSWORD_HASH = pwd_context.hash(ADMIN_PASSWORD)

# Créer base de données
os.makedirs(QR_FOLDER, exist_ok=True)

class Ticket(Base):
    __tablename__ = "tickets"
    ticket_number = Column(Integer, primary_key=True, index=True)
    status = Column(String, default='invalide')
    qr_hash = Column(String, unique=True, nullable=True)
    timestamp = Column(DateTime, default=func.now())

def init_db():
    Base.metadata.create_all(bind=engine)

init_db()

# Interface mobile HTML
with open("static/icon.png", "rb") as f: pass  # vérifie que l'icône existe
MOBILE_TEMPLATE = """
<!DOCTYPE html>
<html lang="fr">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no"/>
  <title>Saint Anne Show</title>
  <link rel="icon" href="/static/icon.png" type="image/png">
  <link rel="apple-touch-icon" href="/static/icon.png">
  <meta name="theme-color" content="#0f172a">
  <meta name="mobile-web-app-capable" content="yes">
  <meta name="apple-mobile-web-app-capable" content="yes">
  <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
  <style>
    * { box-sizing: border-box; }
    html, body {
      margin: 0;
      padding: 0;
      height: 100%;
      width: 100%;
      font-family: 'Segoe UI', Roboto, sans-serif;
      background: linear-gradient(to bottom right, #0f172a, #1e293b);
      color: white;
      overflow-x: hidden;
    }
    .page {
      display: none;
      flex-direction: column;
      align-items: center;
      justify-content: flex-start;
      min-height: 100vh;
      padding: 30px 20px 100px;
    }
    .page.active { display: flex; }

    .logo img {
      width: 220px;
      margin-bottom: 30px;
    }

    input, select {
      padding: 16px;
      margin: 12px 0;
      border: none;
      border-radius: 12px;
      font-size: 16px;
      width: 100%;
      max-width: 360px;
    }

    button {
      padding: 16px;
      margin-bottom: 16px;
      font-size: 17px;
      font-weight: bold;
      border: none;
      border-radius: 12px;
      cursor: pointer;
      width: 100%;
      max-width: 360px;
      color: white;
    }

    .validate { background-color: #22c55e; }
    .verify   { background-color: #3b82f6; }
    .history  { background-color: #f97316; }
    .export   { background-color: #facc15; color: black; }
    .delete   { background-color: #ef4444; }

    .nav-links {
      text-align: center;
      margin-top: 20px;
    }
    .nav-links button {
      background: none;
      border: none;
      color: #a5b4fc;
      font-size: 15px;
      margin: 0 8px;
      text-decoration: underline;
      cursor: pointer;
    }

    #result-validation, #result-verification, #result {
      margin-top: 20px;
      font-size: 16px;
      font-weight: bold;
      color: inherit;
    }

    .input-group {
      display: flex;
      gap: 8px;
      width: 100%;
      max-width: 360px;
    }

    .input-group input {
      flex: 1;
      margin: 12px 0;
    }

    .input-group button {
      width: 60px;
      margin: 12px 0;
      padding: 16px;
      background: #475569;
      border-radius: 12px;
      color: white;
      font-weight: bold;
      font-size: 20px;
    }
  </style>
</head>
<body>

  <!-- Page Validation -->
  <div class="page active" id="validation">
    <div class="logo"><img src="/static/logo.png" alt="Sainte Anne Show"></div>
    <div class="input-group">
      <input type="text" inputmode="numeric" id="ticketInput" placeholder="Ex: 123,456,789" autocomplete="off" pattern="[0-9,]*">
      <button type="button" onclick="insertComma('ticketInput')">,</button>
    </div>
    <button class="validate" onclick="validateTicket()">✅ Valider</button>
    <div id="result-validation"></div>
    <div class="nav-links">
      <button onclick="showPage('verification')">🔍 Vérifier</button>
      <button onclick="showPage('admin')">🛠️ Admin</button>
    </div>
  </div>

  <!-- Page Vérification -->
  <div class="page" id="verification">
    <div class="logo"><img src="/static/logo.png" alt="Sainte Anne Show"></div>
    <div class="input-group">
      <input type="text" inputmode="numeric" id="ticketInputVerify" placeholder="Ex: 123,456,789" autocomplete="off" pattern="[0-9,]*" oninput="formatTicketInput(event)">
      <button type="button" onclick="insertComma('ticketInputVerify')">,</button>
    </div>
    <button class="verify" onclick="verifyTicket()">🔍 Vérifier</button>
    <div id="result-verification"></div>
    <div class="nav-links">
      <button onclick="showPage('validation')">✅ Valider</button>
      <button onclick="showPage('admin')">🛠️ Admin</button>
    </div>
  </div>

  <!-- Page Admin -->
  <div class="page" id="admin">
    <div class="logo"><img src="/static/logo.png" alt="Sainte Anne Show"></div>
    <input type="password" id="adminPass" placeholder="Mot de passe admin">
    <select id="statusFilter">
      <option value="">Tous les statuts</option>
      <option value="validé">Validé</option>
      <option value="invalide">Invalide</option>
    </select>
    <button class="history" onclick="loadHistory()">📄 Historique</button>
    <button class="export" onclick="exportData()">📤 Exporter (.docx)</button>
    <div class="input-group">
      <input type="text" inputmode="numeric" id="deleteTicket" placeholder="Ticket à supprimer (vide = tous)" autocomplete="off" pattern="[0-9,]*" oninput="formatTicketInput(event)">
      <button type="button" onclick="insertComma('deleteTicket')">,</button>
    </div>
    <button class="delete" onclick="deleteValidated()">🗑️ Supprimer</button>
    <div class="nav-links">
      <button onclick="showPage('validation')">✅ Valider</button>
      <button onclick="showPage('verification')">🔍 Vérifier</button>
    </div>
    <div id="result"></div>
    <div id="historyList"></div>
  </div>

 <script>
  function showPage(id) {
    document.querySelectorAll(".page").forEach(p => p.classList.remove("active"));
    document.getElementById(id).classList.add("active");
  }
  
  function insertComma(inputId) {
    const input = document.getElementById(inputId);
    if (input.value && !input.value.endsWith(',')) {
      input.value += ',';
    }
    input.focus();
  }
  
  // Fonction pour formater les numéros de tickets (remplace les points par des virgules)
  function formatTicketInput(event) {
    const input = event.target;
    input.value = input.value.replace(/\./g, ',');
  }
  
  // Fonction utilitaire pour normaliser les numéros de tickets
  function parseTicketNumbers(raw) {
    raw = raw.replace(/\./g, ',');
    return raw.split(',').map(n => n.trim()).filter(n => n.length > 0);
  }

  const apiBase = window.location.origin;

  async function validateTicket() {
    const t = document.getElementById('ticketInput').value;
    if (!t) return alert("Veuillez entrer un numéro de ticket.");
    const r = await fetch(`${apiBase}/validate`, {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ticket: t})
    });
    const d = await r.json();
    const result = document.getElementById('result-validation');
    result.innerText = d.message || d.error;
    result.style.color = d.message ? "#22c55e" : "red";
    document.getElementById('ticketInput').value = '';
    document.getElementById('ticketInput').focus();
  }

  async function verifyTicket() {
    const t = document.getElementById('ticketInputVerify').value;
    if (!t) return alert("Veuillez entrer un numéro de ticket.");
    const r = await fetch(`${apiBase}/verify?ticket=${encodeURIComponent(t)}`);
    const d = await r.json();
    const result = document.getElementById('result-verification');

    if (d.results) {
      result.innerHTML = d.results.map(item => {
        let color;
        if (item.status.includes("validé")) color = "#22c55e";
        else if (item.status.includes("invalide")) color = "#ef4444";
        else color = "#facc15";
        return `<div style="color:${color}">Ticket ${item.ticket}: ${item.status}</div>`;
      }).join('');
    } else {
      result.innerText = d.error || "Erreur inconnue";
      result.style.color = "red";
    }
    document.getElementById('ticketInputVerify').value = '';
    document.getElementById('ticketInputVerify').focus();
  }

  async function exportData() {
    const pwd = document.getElementById('adminPass').value;
    if (!pwd) {
      alert("Veuillez entrer le mot de passe admin.");
      return;
    }
    const r = await fetch(`${apiBase}/export_word`, {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({password: pwd})
    });
    if (r.ok) {
      const blob = await r.blob();
      const a = document.createElement('a');
      a.href = URL.createObjectURL(blob);
      a.download = 'tickets.docx';
      a.click();
    } else {
      alert("Accès refusé - mot de passe incorrect.");
    }
  }

  async function loadHistory() {
    const pwd = document.getElementById('adminPass').value;
    if (!pwd) {
      alert("Veuillez entrer le mot de passe admin.");
      return;
    }
    
    const status = document.getElementById('statusFilter').value;
    const r = await fetch(`${apiBase}/history?status=${status}`, {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({password: pwd})
    });
    const div = document.getElementById('historyList');
    try {
      const data = await r.json();
      if (data.error) {
        div.innerHTML = `<span style='color:red'>Erreur : ${data.error}</span>`;
        return;
      }
      const list = data.results || data;
      if (Array.isArray(list) && list.length > 0) {
        div.innerHTML = list.map(e => `<div>🕒 ${e}</div>`).join('');
      } else if (Array.isArray(list)) {
        div.innerHTML = "<em>Aucun ticket à afficher.</em>";
      } else {
        div.innerHTML = "<em>Aucun ticket à afficher.</em>";
      }
    } catch (err) {
      div.innerHTML = "<span style='color:red'>Erreur de chargement de l'historique.</span>";
    }
  }

  async function deleteValidated() {
    const pwd = document.getElementById('adminPass').value;
    const ticket = document.getElementById('deleteTicket').value;
    
    if (!pwd) {
      alert("Veuillez entrer le mot de passe admin.");
      return;
    }
    
    const confirmDelete = confirm(ticket ? `Supprimer les tickets validés N°${ticket} ?` : "Confirmer la suppression de tous les tickets validés ?");
    if (!confirmDelete) return;

    const r = await fetch(`${apiBase}/delete_validated`, {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({password: pwd, ticket: ticket || null})
    });
    const d = await r.json();
    const resultDiv = document.getElementById('result');
    resultDiv.innerText = d.message || d.error;
    resultDiv.style.color = d.message ? "#22c55e" : "red";
    document.getElementById('deleteTicket').value = '';
  }
</script>
</body>
</html>
"""
# Flask App
app = Flask(__name__)
CORS(app)

@app.after_request
def add_headers(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return response

@app.route('/')
def home():
    return render_template_string(MOBILE_TEMPLATE)

@app.route('/validate', methods=['POST'])
def validate():
    try:
        data = request.get_json()
        raw = str(data.get('ticket', ''))
        raw = raw.replace('.', ',')
        numbers = [n.strip() for n in raw.split(',') if n.strip()]
        if not numbers:
            return jsonify({"error": "Aucun numéro de ticket fourni"}), 400

        messages = []
        errors = []
        db = SessionLocal()
        for n in numbers:
            if not n.isdigit():
                errors.append(f"'{n}' invalide")
                continue
            t = int(n)
            ticket = db.query(Ticket).filter_by(ticket_number=t).first()
            if ticket:
                ticket.status = f"validé - {t}"
            else:
                ticket = Ticket(ticket_number=t, status=f"validé - {t}")
                db.add(ticket)
            messages.append(f"Ticket {t} validé")
        db.commit()
        db.close()

        response = {}
        if messages:
            response["message"] = " | ".join(messages)
        if errors:
            response["error"] = " | ".join(errors)
        return jsonify(response)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/verify')
def verify():
    try:
        raw = request.args.get('ticket', '')
        raw = raw.replace('.', ',')
        numbers = [n.strip() for n in raw.split(',') if n.strip()]
        if not numbers:
            return jsonify({"error": "Aucun numéro de ticket fourni"}), 400

        results = []
        with SessionLocal() as db:
            for n in numbers:
                if not n.isdigit():
                    results.append({"ticket": n, "status": "numéro invalide"})
                    continue
                t = int(n)
                ticket = db.query(Ticket).filter_by(ticket_number=t).first()
                if ticket:
                    results.append({"ticket": t, "status": ticket.status})
                else:
                    ticket = Ticket(ticket_number=t, status=f"invalide - {t}")
                    db.add(ticket)
                    db.commit()
                    results.append({"ticket": t, "status": ticket.status})

        return jsonify({"results": results})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/export_word', methods=['POST'])
def export_word():
    try:
        data = request.get_json()
        if not pwd_context.verify(data.get('password', ''), ADMIN_PASSWORD_HASH):
            return jsonify({"error": "Accès refusé"}), 401

        doc = Document()
        doc.add_heading("Tickets Validés", 0)

        db = SessionLocal()
        results = db.query(Ticket).filter(
          or_(
        Ticket.status == 'validé',
        Ticket.status.like('validé%')
       )
    ).all()
        db.close()

        if not results:
            doc.add_paragraph("Aucun ticket validé.")
        else:
            for ticket in results:
                doc.add_paragraph(f"Ticket {ticket.ticket_number} - Validé le {ticket.timestamp}")

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='tickets_valides.docx'
        )
    except Exception as e:
        return jsonify({"error": f"Erreur export: {str(e)}"}), 500

@app.route('/admin', methods=['POST'])
def admin():
    try:
        data = request.get_json()
        if pwd_context.verify(data.get('password', ''), ADMIN_PASSWORD_HASH):
            return jsonify({"success": True})
        return jsonify({"success": False}), 401
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/delete_validated', methods=['POST'])
def delete_validated():
    try:
        data = request.get_json()
        if not pwd_context.verify(data.get('password', ''), ADMIN_PASSWORD_HASH):
            return jsonify({"error": "Accès refusé"}), 401

        ticket_input = data.get("ticket", "")
        db = SessionLocal()
        
        # Si ticket_input est vide ou None, supprimer tous les tickets validés
        if not ticket_input:
            deleted = db.query(Ticket).filter(
                or_(
                    Ticket.status == "validé",
                    Ticket.status.like("validé%")
                )
            ).delete()
        else:
            # Sinon, traiter les numéros (peut être plusieurs séparés par des virgules)
            raw = str(ticket_input).replace('.', ',')
            numbers = [n.strip() for n in raw.split(',') if n.strip()]
            valid_numbers = [int(n) for n in numbers if n.isdigit()]
            
            if valid_numbers:
                deleted = db.query(Ticket).filter(
                    and_(
                        Ticket.ticket_number.in_(valid_numbers),
                        or_(
                            Ticket.status == "validé",
                            Ticket.status.like("validé%")
                        )
                    )
                ).delete()
            else:
                deleted = 0
                                        
        db.commit()
        db.close()
        return jsonify({"message": f"{deleted} ticket(s) supprimé(s)."})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/history', methods=['POST', 'GET'])
def history():
    try:
        # Vérifier le mot de passe s'il y a une requête POST
        if request.method == 'POST':
            data = request.get_json()
            if not pwd_context.verify(data.get('password', ''), ADMIN_PASSWORD_HASH):
                return jsonify({"error": "Accès refusé"}), 401
        
        status = request.args.get("status")
        db = SessionLocal()
        query = db.query(Ticket)

        if status in ("validé", "invalide"):
            query = query.filter(Ticket.status.like(f"{status}%"))

        results = query.order_by(Ticket.timestamp.desc()).limit(MAX_HISTORY_ENTRIES).all()
        db.close()

        return jsonify({
            "results": [
                f"Ticket {r.ticket_number} - {r.status} - {r.timestamp}" for r in results
            ]
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/ping')
def ping():
    return "pong", 200
if __name__ == '__main__':
 app.run(host='0.0.0.0', port=FLASK_PORT)
